from pathlib import Path

from app.models.entities import MarkdownDocument


class DocumentPipeline:
    def __init__(self, markdown_dir: Path):
        self.markdown_dir = Path(markdown_dir)

    def convert_to_markdown(self, source_path: Path) -> MarkdownDocument:
        source = Path(source_path)
        suffix = source.suffix.lower()
        converters = {
            ".txt": self._convert_text,
            ".md": self._convert_markdown,
            ".docx": self._convert_docx,
            ".pdf": self._convert_pdf,
            ".xlsx": self._convert_xlsx,
        }
        if suffix not in converters:
            raise ValueError(f"Unsupported document type: {suffix}")

        content = converters[suffix](source).strip()
        markdown_path = self.markdown_dir / f"{source.stem}.md"
        self.markdown_dir.mkdir(parents=True, exist_ok=True)
        markdown_path.write_text(content, encoding="utf-8")
        return MarkdownDocument(source_path=source, markdown_path=markdown_path, content=content)

    def _convert_text(self, source: Path) -> str:
        return f"# {source.stem}\n\n{source.read_text(encoding='utf-8').strip()}"

    def _convert_markdown(self, source: Path) -> str:
        return source.read_text(encoding="utf-8")

    def _convert_docx(self, source: Path) -> str:
        try:
            import docx
        except ImportError as exc:
            raise RuntimeError("python-docx is required to parse .docx files") from exc
        document = docx.Document(str(source))
        lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        return f"# {source.stem}\n\n" + "\n\n".join(lines)

    def _convert_pdf(self, source: Path) -> str:
        try:
            import fitz
        except ImportError as exc:
            raise RuntimeError("PyMuPDF is required to parse .pdf files") from exc
        parts = []
        with fitz.open(str(source)) as document:
            for page in document:
                text = page.get_text().strip()
                if text:
                    parts.append(text)
        return f"# {source.stem}\n\n" + "\n\n".join(parts)

    def _convert_xlsx(self, source: Path) -> str:
        try:
            import openpyxl
        except ImportError as exc:
            raise RuntimeError("openpyxl is required to parse .xlsx files") from exc
        workbook = openpyxl.load_workbook(source, read_only=True, data_only=True)
        sections = [f"# {source.stem}"]
        for sheet in workbook.worksheets:
            sections.append(f"## {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                values = [str(value) for value in row if value is not None]
                if values:
                    sections.append("| " + " | ".join(values) + " |")
        return "\n\n".join(sections)
